#!/usr/bin/env python3
"""
Explore the TC sheet and TN.doc files to understand the data structure.
"""

import openpyxl
from docx import Document
import os

# Load the Excel file and get TC sheet
print("=" * 60)
print("EXPLORING TC SHEET")
print("=" * 60)
wb = openpyxl.load_workbook('offense_codes_updated.xlsx')
ws = wb['TC']

print(f"Total rows in TC sheet: {ws.max_row}")
print("\nFirst 15 rows:")
for row_num, row in enumerate(ws.iter_rows(min_row=1, max_row=15, max_col=6, values_only=True), 1):
    print(f"Row {row_num}: {row}")

# Get a few sample citations from column B
print("\n" + "=" * 60)
print("SAMPLE CITATIONS FROM COLUMN B")
print("=" * 60)
citations = []
for row in ws.iter_rows(min_row=2, max_row=20, min_col=2, max_col=2, values_only=True):
    if row[0]:
        citations.append(row[0])
print("Sample citations:", citations[:10])

# Explore a TN.doc file
print("\n" + "=" * 60)
print("EXPLORING TN.DOC FILE (tn.545.docx)")
print("=" * 60)
doc_path = 'TN.doc/tn.545.docx'
if os.path.exists(doc_path):
    doc = Document(doc_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("\nFirst 25 paragraphs with content:")
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"[{i}] {text[:200]}..." if len(text) > 200 else f"[{i}] {text}")
            count += 1
            if count >= 25:
                break

# List all TN.doc files
print("\n" + "=" * 60)
print("TN.DOC FILES (first 20)")
print("=" * 60)
tn_files = os.listdir('TN.doc')
print(f"Total files: {len(tn_files)}")
print("Sample filenames:", sorted(tn_files)[:20])

wb.close()
